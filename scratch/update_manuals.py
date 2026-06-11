import docx
import os

docs_dir = r"C:\Development\Antigravity\RRHH_Objetivos\docs\Manuales"
admin_doc_path = os.path.join(docs_dir, "Manual_Administrador.docx")
emp_doc_path = os.path.join(docs_dir, "Manual_Empleado.docx")

def add_appendix(doc_path, text_blocks):
    if not os.path.exists(doc_path):
        print(f"File not found: {doc_path}")
        return
        
    doc = docx.Document(doc_path)
    
    doc.add_page_break()
    doc.add_heading('Anexo: Nuevas Funcionalidades', level=1)
    
    for title, body in text_blocks:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)
        
    doc.save(doc_path)
    print(f"Updated {doc_path}")

admin_blocks = [
    ("Objetivos y Seguimiento", "Como Administrador, notarás que la vista de Objetivos se divide en solapas específicas para facilitar la gestión:\n- Objetivos de mi Equipo: Aquí verás únicamente a los colaboradores que te reportan de forma directa, permitiéndote un seguimiento enfocado como líder.\n- Toda la Organización: Esta solapa es exclusiva de tu perfil administrador y te permite tener una visión global y auditar el progreso de cualquier colaborador de la empresa."),
    ("Autoevaluación", "Además de poder visualizar las autoevaluaciones de tu equipo, como líder también tienes el deber de completar tus propias autoevaluaciones. Para ello, dirígete a la solapa 'Pendientes de Autoevaluar', donde encontrarás tus propios objetivos listos para que ingreses tus comentarios y evidencias."),
    ("Calendario", "El Calendario está diseñado para darte visibilidad sobre los Hitos del Proceso (ej. Fechas de corte para evaluaciones de mitad de año y cierre final) a nivel compañía, sin sobrecargarse visualmente con las tareas operativas pendientes de cada empleado individual."),
    ("Progreso del Objetivo", "La columna Progreso te permite saber en qué etapa del proceso de evaluación se encuentra cada objetivo. A medida que completes las distintas fases, la barra avanzará automáticamente:\n- 0%: El objetivo está activo.\n- 33%: Has completado y enviado tu Autoevaluación exitosamente.\n- 66%: Tu líder ha completado la instancia de Feedback de Mitad de Año.\n- 100%: El ciclo ha cerrado con la Evaluación Final.")
]

emp_blocks = [
    ("Progreso del Objetivo", "La columna Progreso te permite saber en qué etapa del proceso de evaluación se encuentra cada objetivo. A medida que completes las distintas fases, la barra avanzará automáticamente:\n- 0%: El objetivo está activo.\n- 33%: Has completado y enviado tu Autoevaluación exitosamente.\n- 66%: Tu líder ha completado la instancia de Feedback de Mitad de Año.\n- 100%: El ciclo ha cerrado con la Evaluación Final.")
]

add_appendix(admin_doc_path, admin_blocks)
add_appendix(emp_doc_path, emp_blocks)
