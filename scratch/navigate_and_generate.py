import os
import sys
import time
import subprocess
import requests
from playwright.sync_api import sync_playwright
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- Configuration Constants ---
PORT = 5118
BASE_URL = f"http://localhost:{PORT}"
SCREENSHOT_DIR = "docs/screenshots"
ADMIN_EMAIL = "ptripodi@permaquim.com"
ADMIN_PASS = "18"
EMP_EMAIL = "ncaldiroli@permaquim.com"
EMP_PASS = "560"

# --- Visual Colors for Docx styling (PQ Blue theme) ---
COLOR_PRIMARY = RGBColor(24, 76, 120)     # Deep PQ Blue (#184C78)
COLOR_SECONDARY = RGBColor(46, 117, 182) # Secondary Blue (#2E75B6)
COLOR_MUTED = RGBColor(100, 100, 100)     # Grey (#646464)
HEX_CALLOUT_BORDER = "184C78"
HEX_CALLOUT_BG = "F2F5F8"

def create_callout_box(doc, text):
    """Creates a beautifully styled callout box for business rules in Word."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Inches(5.8)
    
    cell = tbl.cell(0, 0)
    # Set background color
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), HEX_CALLOUT_BG)
    tcPr.append(shd)
    
    # Set left border thick, others none
    tcBorders = OxmlElement('w:tcBorders')
    
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '36')  # 4.5 pt
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), HEX_CALLOUT_BORDER)
    tcBorders.append(left)
    
    for side in ['top', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
    tcPr.append(tcBorders)
    
    # Add text
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.right_indent = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    
    run_label = p.add_run("Regla de Negocio: ")
    run_label.font.name = 'Calibri'
    run_label.font.size = Pt(10)
    run_label.bold = True
    run_label.font.color.rgb = COLOR_PRIMARY
    
    run_text = p.add_run(text)
    run_text.font.name = 'Calibri'
    run_text.font.size = Pt(10)
    run_text.italic = True
    run_text.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = COLOR_SECONDARY
    return p

def add_body_text(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_screenshot(doc, img_name, caption):
    """Inserts a screenshot centered with a caption."""
    path = os.path.join(SCREENSHOT_DIR, img_name)
    if not os.path.exists(path):
        print(f"Warning: Screenshot {path} not found! Skipping image insert.")
        return
    
    # Centered image
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(4)
    
    run_img = p_img.add_run()
    run_img.add_picture(path, width=Inches(5.5))
    
    # Centered caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(0)
    p_cap.paragraph_format.space_after = Pt(12)
    
    run_cap = p_cap.add_run(f"Ilustración: {caption}")
    run_cap.font.name = 'Calibri'
    run_cap.font.size = Pt(9.5)
    run_cap.italic = True
    run_cap.font.color.rgb = COLOR_MUTED

def add_header_footer(doc, title_text):
    """Sets custom header and footer for document sections."""
    for section in doc.sections:
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run(f"{title_text} | PQ-Talent Portal")
        hrun.font.name = 'Calibri'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = COLOR_MUTED
        
        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Confidencial - Uso Interno Permaquim S.A.")
        frun.font.name = 'Calibri'
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = COLOR_MUTED

def generate_cover_page(doc, title, subtitle, author):
    """Creates a beautiful corporate cover page."""
    # Top spacing
    for _ in range(3):
        doc.add_paragraph()
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run(title)
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(28)
    run_title.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run(subtitle)
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(14)
    run_sub.italic = True
    run_sub.font.color.rgb = COLOR_SECONDARY
    
    # Decorative line
    p_line = doc.add_paragraph()
    p_line_run = p_line.add_run("━" * 45)
    p_line_run.bold = True
    p_line_run.font.color.rgb = COLOR_PRIMARY
    
    for _ in range(8):
        doc.add_paragraph()
        
    p_meta = doc.add_paragraph()
    run_meta = p_meta.add_run(f"Preparado por: {author}\nFecha: Mayo 2026\nVersión: 1.0\nEmpresa: Permaquim S.A.")
    run_meta.font.name = 'Calibri'
    run_meta.font.size = Pt(10)
    run_meta.font.color.rgb = COLOR_MUTED
    
    doc.add_page_break()

# --- Main Script Execution ---
def main():
    os.makedirs(SCREENSHOT_DIR, exist_ok=True)
    os.makedirs("docs", exist_ok=True)

    # 1. Start Blazor Server in the background
    server_process = None
    server_running = False
    
    print("Checking if port 5118 is already in use...")
    try:
        r = requests.get(f"{BASE_URL}/login", timeout=2)
        if r.status_code == 200:
            print("Server is already running on port 5118. Reusing existing instance.")
            server_running = True
    except Exception:
        pass
        
    if not server_running:
        print("Starting Blazor Server local instance in background...")
        # Start dotnet run
        server_process = subprocess.Popen(
            ["dotnet", "run", "--project", "Objetivos.Web/Objetivos.Web.csproj", "--urls", BASE_URL],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True
        )
        
        # Wait up to 30 seconds for server availability
        for i in range(30):
            try:
                r = requests.get(f"{BASE_URL}/login", timeout=2)
                if r.status_code == 200:
                    print("Blazor Server launched successfully and is listening on port 5118!")
                    server_running = True
                    break
            except Exception:
                pass
            time.sleep(1)
            
        if not server_running:
            print("CRITICAL: Blazor Server failed to start on port 5118. Exiting.")
            if server_process:
                server_process.kill()
            sys.exit(1)

    # 2. Run Playwright automation to take screenshots
    try:
        run_playwright_scenarios()
    except Exception as ex:
        print(f"Error during Playwright execution: {ex}")
    finally:
        # Stop local background server if we started it
        if server_process:
            print("Stopping background Blazor Server process...")
            server_process.terminate()
            try:
                server_process.wait(timeout=5)
            except subprocess.TimeoutExpired:
                server_process.kill()
            print("Server process terminated.")

    # 3. Generate Word Documents
    print("Generating manuals...")
    generate_admin_manual()
    generate_employee_manual()
    print("Done! Both manuals compiled successfully under the 'docs' directory.")

def run_playwright_scenarios():
    print("Launching browser via Playwright...")
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        # 1440x900 viewport for clean presentation of Radzen Dashboard charts
        context = browser.new_context(viewport={"width": 1440, "height": 900})
        page = context.new_page()
        
        # Set generous timeouts
        page.set_default_navigation_timeout(60000)
        page.set_default_timeout(60000)

        # Block external fonts to avoid hanging in offline environment
        def handle_route(route):
            url = route.request.url
            if "fonts.googleapis" in url or "fonts.gstatic" in url or url.endswith((".woff", ".woff2", ".ttf")):
                route.abort()
            else:
                route.continue_()
        context.route("**/*", handle_route)

        # Helper function for safe navigation
        def safe_goto(url, name, wait_time=2000):
            print(f"Navigating to {name} ({url})...")
            try:
                page.goto(url, timeout=30000, wait_until="domcontentloaded")
                page.wait_for_timeout(wait_time)
                # Wait for any Blazor loading progress bars to disappear
                try:
                    page.wait_for_selector(".rz-progressbar", state="hidden", timeout=3000)
                except Exception:
                    pass
                return True
            except Exception as e:
                print(f"WARNING: Failed to load {name} page: {e}")
                return False

        # Helper function to navigate using the Radzen left menu (User Rules 3 & 4 verification)
        def safe_navigate_menu(menu_path, expected_url, name, wait_time=2000):
            """
            menu_path: list of menu items to click in sequence, e.g. ["Administración", "Configuración"] or ["Cursos"]
            expected_url: the expected URL path (e.g. "/admin/configuracion")
            """
            print(f"Navigating via menu to {name} ({menu_path}) -> {expected_url}...")
            try:
                for text in menu_path:
                    # Look for the Radzen panel menu item with the exact text
                    locator = page.locator(".rz-navigation-item-link, .rz-panel-menu-link, a").filter(has_text=text).first
                    if not locator.is_visible():
                        locator = page.locator(f'text="{text}"').first
                    
                    # Use force=True to prevent menu expanding transition from blocking click
                    locator.click(timeout=10000, force=True)
                    page.wait_for_timeout(1000)
                
                # Verify we reached the expected URL or path
                page.wait_for_url(f"**{expected_url}", timeout=15000)
                page.wait_for_timeout(wait_time)
                
                # Wait for any Blazor loading progress bars to disappear
                try:
                    page.wait_for_selector(".rz-progressbar", state="hidden", timeout=3000)
                except Exception:
                    pass
                return True
            except Exception as e:
                print(f"WARNING: Failed to navigate via menu to {name}: {e}")
                # Fallback to direct navigation
                print(f"Falling back to direct navigation for {name} ({expected_url})...")
                return safe_goto(f"{BASE_URL}{expected_url}", name, wait_time)

        # Helper function for safe click
        def safe_click(selector, name, wait_time=1500):
            print(f"Clicking {name}...")
            try:
                page.click(selector, timeout=15000)
                page.wait_for_timeout(wait_time)
                return True
            except Exception as e:
                print(f"WARNING: Failed to click {name}: {e}")
                return False

        # Helper function for screenshot
        def safe_screenshot(path, name):
            try:
                page.screenshot(path=path, timeout=5000)
                print(f"Captured: {os.path.basename(path)}")
            except Exception as e:
                print(f"WARNING: Failed to capture screenshot {name}: {e}")

        # =========================================================================
        # SCENARIO A: ADMINISTRADOR (PABLO TRIPODI)
        # =========================================================================
        print("--- Running Scenario: Administrador (Pablo Tripodi) ---")
        
        # 1. Login Page
        if safe_goto(f"{BASE_URL}/login", "Login Page", 1000):
            safe_screenshot(f"{SCREENSHOT_DIR}/01_admin_login.png", "admin_login")
            
            # Fill credentials
            print("Entering credentials for Admin...")
            try:
                page.fill('input[name="Email"]', ADMIN_EMAIL)
                page.fill('input[name="Password"]', ADMIN_PASS)
                page.click("button.login-btn")
            except Exception as e:
                print(f"WARNING: Login fields or button not found: {e}")
            
            # 2. Dashboard - Team tab (Default)
            print("Waiting for dashboard redirect...")
            try:
                page.wait_for_url(f"{BASE_URL}/dashboard", timeout=20000)
                page.wait_for_timeout(3000) # wait for chart animations to finish
                safe_screenshot(f"{SCREENSHOT_DIR}/02_admin_dashboard_team.png", "admin_dashboard_team")
            except Exception as e:
                print(f"WARNING: Dashboard redirection failed: {e}")
            
            # 3. Dashboard - Personal tab
            # Click the tab "Mis Objetivos"
            if safe_click('text="Mis Objetivos"', "Mis Objetivos Tab", 1500):
                safe_screenshot(f"{SCREENSHOT_DIR}/03_admin_dashboard_personal.png", "admin_dashboard_personal")
            
            # 4. Objetivos y Competencias Page (Navigated via Left Menu)
            if safe_navigate_menu(["Objetivos y Competencias"], "/objetivos", "Objetivos y Competencias", 2000):
                safe_screenshot(f"{SCREENSHOT_DIR}/04_admin_objetivos.png", "admin_objetivos")
            
            # 5. Seguimientos Page (Navigated via Left Menu)
            if safe_navigate_menu(["Seguimientos"], "/seguimientos", "Seguimientos", 2000):
                safe_screenshot(f"{SCREENSHOT_DIR}/05_admin_seguimientos.png", "admin_seguimientos")
            
            # 6. Colaborador Detalle Page (Nicolas Caldiroli ID 2)
            print("Navigating to Colaborador Detalle...")
            caldiroli_clicked = False
            try:
                # Click the "Ver Detalle" button for Nicolas Caldiroli in the table
                caldiroli_row = page.locator("tr", has_text="Caldiroli")
                if caldiroli_row.count() > 0:
                    caldiroli_row.locator('text="Ver Detalle"').click(timeout=5000)
                    page.wait_for_timeout(2500)
                    caldiroli_clicked = True
            except Exception as e:
                print(f"WARNING: Click 'Ver Detalle' for Caldiroli failed: {e}")
            
            if not caldiroli_clicked:
                safe_goto(f"{BASE_URL}/seguimientos/2", "Colaborador Detalle", 2500)
            else:
                safe_screenshot(f"{SCREENSHOT_DIR}/06_admin_empleado_detalle.png", "admin_empleado_detalle")
            
            # 7. Objetivo Detalle Page (Nicolas's Support Ticket Objective 1)
            print("Navigating to Objetivo Detalle...")
            objective_clicked = False
            try:
                # Click the first "Ver" button in the objectives grid
                page.locator(".rz-data-grid table").locator('text="Ver"').first.click(timeout=5000)
                page.wait_for_timeout(2500)
                objective_clicked = True
            except Exception as e:
                print(f"WARNING: Click 'Ver' objective failed: {e}")
                
            if not objective_clicked:
                safe_goto(f"{BASE_URL}/seguimientos/2/objetivo/1", "Objetivo Detalle", 2500)
            else:
                safe_screenshot(f"{SCREENSHOT_DIR}/07_admin_objetivo_detalle.png", "admin_objetivo_detalle")
            
            # 8. Autoevaluaciones Page (Navigated via Left Menu)
            if safe_navigate_menu(["Autoevaluación"], "/autoevaluacion", "Autoevaluaciones", 2000):
                safe_screenshot(f"{SCREENSHOT_DIR}/08_admin_autoevaluacion.png", "admin_autoevaluacion")
            
            # 9. Evaluaciones Page (Navigated via Left Menu)
            if safe_navigate_menu(["Evaluación"], "/evaluacion", "Evaluaciones", 2000):
                safe_screenshot(f"{SCREENSHOT_DIR}/09_admin_evaluacion.png", "admin_evaluacion")
            
            # 10. Cursos Page (Navigated via Left Menu)
            if safe_navigate_menu(["Cursos"], "/cursos", "Cursos", 2000):
                safe_screenshot(f"{SCREENSHOT_DIR}/10_admin_cursos.png", "admin_cursos")
            
            # 11. Calendario Page (Navigated via Left Menu)
            if safe_navigate_menu(["Calendario"], "/calendario", "Calendario", 2000):
                safe_screenshot(f"{SCREENSHOT_DIR}/11_admin_calendario.png", "admin_calendario")
            
            # 12. Guía Page (Navigated via Left Menu)
            if safe_navigate_menu(["Guía"], "/guia", "Guía", 2000):
                safe_screenshot(f"{SCREENSHOT_DIR}/12_admin_guia.png", "admin_guia")
            
            # 13. Configuración Page (Navigated via Left Menu - nested under Administración)
            if safe_navigate_menu(["Administración", "Configuración"], "/admin/configuracion", "Configuracion", 2500):
                safe_screenshot(f"{SCREENSHOT_DIR}/13_admin_configuracion.png", "admin_configuracion")
            
            # 14. Usuarios Page (Navigated via Left Menu - nested under Administración)
            if safe_navigate_menu(["Administración", "Usuarios"], "/admin/usuarios", "Usuarios", 2500):
                safe_screenshot(f"{SCREENSHOT_DIR}/14_admin_usuarios.png", "admin_usuarios")
            
            # Logout
            if safe_click('text="Cerrar Sesión"', "Cerrar Sesión", 1500):
                try:
                    page.wait_for_url(f"{BASE_URL}/login", timeout=10000)
                    print("Logged out from Admin account.")
                except Exception:
                    print("WARNING: Redirect to login after logout timed out.")

        # =========================================================================
        # SCENARIO B: EMPLEADO (NICOLAS CALDIROLI)
        # =========================================================================
        print("--- Running Scenario: Empleado (Nicolas Caldiroli) ---")
        
        # 1. Login Page
        if "/login" in page.url or safe_goto(f"{BASE_URL}/login", "Login Page", 1000):
            page.wait_for_timeout(1000)
            safe_screenshot(f"{SCREENSHOT_DIR}/01_empleado_login.png", "empleado_login")
            
            # Fill credentials
            print("Entering credentials for Empleado...")
            try:
                page.fill('input[name="Email"]', EMP_EMAIL)
                page.fill('input[name="Password"]', EMP_PASS)
                page.click("button.login-btn")
            except Exception as e:
                print(f"WARNING: Login fields or button not found: {e}")
            
            # 2. Dashboard - Personal (Only tab shown for standard employees)
            print("Waiting for dashboard redirect...")
            try:
                page.wait_for_url(f"{BASE_URL}/dashboard", timeout=20000)
                page.wait_for_timeout(3000)
                safe_screenshot(f"{SCREENSHOT_DIR}/02_empleado_dashboard.png", "empleado_dashboard")
            except Exception as e:
                print(f"WARNING: Dashboard redirection failed: {e}")
            
            # 3. Objetivos y Competencias (Personal - Navigated via Left Menu)
            if safe_navigate_menu(["Objetivos y Competencias"], "/objetivos", "Objetivos y Competencias", 2000):
                safe_screenshot(f"{SCREENSHOT_DIR}/03_empleado_objetivos.png", "empleado_objetivos")
            
            # 4. Objetivo Detalle (Personal view - opened via clicking the grid row action)
            print("Navigating to Empleado Objetivo Detalle...")
            emp_objective_clicked = False
            try:
                page.locator(".rz-data-grid table").locator("button").first.click(timeout=5000)
                page.wait_for_timeout(2500)
                emp_objective_clicked = True
            except Exception as e:
                print(f"WARNING: Click personal objective 'Ver' failed: {e}")
                
            if not emp_objective_clicked:
                safe_goto(f"{BASE_URL}/seguimientos/2/objetivo/1", "Objetivo Detalle", 2500)
            else:
                safe_screenshot(f"{SCREENSHOT_DIR}/04_empleado_objetivo_detalle.png", "empleado_objetivo_detalle")
            
            # 5. Autoevaluaciones Page (Personal - Navigated via Left Menu)
            if safe_navigate_menu(["Autoevaluación"], "/autoevaluacion", "Autoevaluación", 2000):
                safe_screenshot(f"{SCREENSHOT_DIR}/05_empleado_autoevaluacion.png", "empleado_autoevaluacion")
            
            # 6. Cursos Page (Personal - Navigated via Left Menu)
            if safe_navigate_menu(["Cursos"], "/cursos", "Cursos", 2000):
                safe_screenshot(f"{SCREENSHOT_DIR}/06_empleado_cursos.png", "empleado_cursos")
            
            # 7. Calendario Page (Personal - Navigated via Left Menu)
            if safe_navigate_menu(["Calendario"], "/calendario", "Calendario", 2000):
                safe_screenshot(f"{SCREENSHOT_DIR}/07_empleado_calendario.png", "empleado_calendario")
            
            # 8. Guía Page (Personal - Navigated via Left Menu)
            if safe_navigate_menu(["Guía"], "/guia", "Guía", 2000):
                safe_screenshot(f"{SCREENSHOT_DIR}/08_empleado_guia.png", "empleado_guia")
            
            # Logout
            if safe_click('text="Cerrar Sesión"', "Cerrar Sesión", 1500):
                try:
                    page.wait_for_url(f"{BASE_URL}/login", timeout=10000)
                    print("Logged out from Empleado account.")
                except Exception:
                    print("WARNING: Redirect to login after logout timed out.")
        
        browser.close()

def generate_admin_manual():
    doc = docx.Document()
    generate_cover_page(doc, "Manual de Usuario - Administrador", "Guía Completa para Jefes, Gerentes y Administradores de Plataforma", "Antigravity AI Agent")
    add_header_footer(doc, "Manual de Administrador")
    
    add_heading_1(doc, "1. Introducción a PQ-Talent (Admin)")
    add_body_text(
        doc,
        "Bienvenido al manual del Administrador y Líder para la plataforma PQ-Talent. "
        "Este portal centraliza la gestión del rendimiento de Permaquim S.A., permitiendo a los líderes "
        "configurar parámetros del sistema, gestionar la nómina de colaboradores directos, definir objetivos estratégicos, "
        "hacer revisiones de mitad de año (Bitácora de seguimiento), responder chats y realizar la evaluación de desempeño final."
    )
    
    add_heading_1(doc, "2. Inicio de Sesión y Acceso")
    add_body_text(
        doc,
        "Para ingresar al portal, navegue al sitio de PQ-Talent. El sistema le presentará la pantalla de inicio de sesión "
        "donde deberá indicar su correo electrónico institucional como usuario y su legajo como contraseña temporal la primera vez. "
        "Al ingresar por primera vez, el sistema solicitará que cambie su contraseña por una que cumpla con los requisitos mínimos de seguridad."
    )
    add_screenshot(doc, "01_admin_login.png", "Pantalla de Inicio de Sesión (Login) para Administrador.")

    create_callout_box(
        doc,
        "Las contraseñas iniciales corresponden exactamente al legajo del usuario. "
        "La contraseña definitiva debe poseer al menos 8 caracteres, conteniendo letras y números combinados. "
        "En caso de olvido, la opción '¿Olvidó su contraseña?' permite generar una temporal e iniciar el flujo de reseteo."
    )

    add_heading_1(doc, "3. Módulo 1: Dashboard Corporativo y de Equipo")
    add_body_text(
        doc,
        "Al autenticarse como Administrador o Jefe con personal a cargo, se le presentará el Dashboard general. "
        "Esta vista cuenta con pestañas diferenciadas para consolidar la visión del equipo frente a los objetivos individuales."
    )
    
    add_heading_2(doc, "3.1 Pestaña: Mi Equipo / Área")
    add_body_text(
        doc,
        "Muestra de forma gráfica la distribución de los objetivos de sus colaboradores directos. "
        "Consta de cuatro tarjetas principales de KPIs (Total Objetivos, En Curso, Vencen Pronto y Pendientes de Revisión), "
        "un gráfico de dona de distribución y un panel lateral derecho con próximos vencimientos y hitos del calendario."
    )
    add_screenshot(doc, "02_admin_dashboard_team.png", "Vista del Dashboard del Equipo - Pestaña Mi Equipo / Área.")

    create_callout_box(
        doc,
        "El estado 'Vencen Pronto' representa un subconjunto de objetivos en estado 'En Curso' o 'Activo' "
        "cuya fecha límite está dentro del margen de alerta establecido en la configuración global (por defecto: 7 días). "
        "Al no ser un estado excluyente, no aparece en la gráfica de torta, sino como una alerta rápida en la tarjeta de KPI superior."
    )

    add_heading_2(doc, "3.2 Pestaña: Mis Objetivos")
    add_body_text(
        doc,
        "Permite alternar para visualizar exclusivamente su rendimiento personal, sus metas personales fijadas por la dirección, "
        "y sus vencimientos personales del periodo."
    )
    add_screenshot(doc, "03_admin_dashboard_personal.png", "Vista del Dashboard Personal del Administrador.")

    add_heading_1(doc, "4. Módulo 2: Objetivos y Competencias")
    add_body_text(
        doc,
        "Este panel permite a los líderes dar seguimiento a la carga de objetivos y competencias de todo su equipo. "
        "Los objetivos se asocian a un Pilar Estratégico Organizacional y a dos Competencias Blandas (Soft Skills) "
        "del diccionario institucional corporativo de Permaquim."
    )
    add_screenshot(doc, "04_admin_objetivos.png", "Panel de Objetivos y Competencias con filtros por historial.")
    
    create_callout_box(
        doc,
        "El sistema valida que la suma de porcentajes de peso de los objetivos de cada colaborador sea exactamente el 100% "
        "para que el ciclo de evaluación del año sea calificado como válido. En caso contrario, se mostrarán alertas informativas."
    )

    add_heading_1(doc, "5. Módulo 3: Seguimientos y Detalle de Colaborador")
    add_body_text(
        doc,
        "En la sección 'Seguimientos' se visualiza una lista de todos los colaboradores a cargo directo. "
        "Se incluye el puesto de trabajo, un semáforo rápido indicador de Rendimiento General (rojo, amarillo, verde) y "
        "el conteo de mensajes sin leer específicos de cada persona."
    )
    add_screenshot(doc, "05_admin_seguimientos.png", "Grilla general de Seguimientos de colaboradores directos.")

    add_heading_2(doc, "5.1 Vista Detallada de Colaborador")
    add_body_text(
        doc,
        "Al hacer clic en 'Ver Detalle' de un colaborador, el líder accede a su ficha técnica. "
        "Esta pantalla expone los datos de contacto, la fecha de ingreso, un gráfico de rendimiento promedio agrupado por "
        "pilar estratégico en forma de columnas, y la lista completa de sus objetivos individuales asignados."
    )
    add_screenshot(doc, "06_admin_empleado_detalle.png", "Ficha técnica de colaborador con gráfico por pilar y grilla de objetivos.")

    add_heading_2(doc, "5.2 Detalle de Objetivo, Bitácora e Interacción (Chat)")
    add_body_text(
        doc,
        "Al ingresar al detalle de un objetivo en particular, el líder accede al canal interactivo. "
        "Aquí puede auditar y gestionar las evidencias de la Bitácora de seguimiento, ver el gráfico de evolución temporal y "
        "chatear de forma directa y bidireccional con el colaborador."
    )
    add_screenshot(doc, "07_admin_objetivo_detalle.png", "Detalle de Objetivo con Bitácora editable y canal de chat bidireccional.")

    create_callout_box(
        doc,
        "El ciclo de vida de una entrada en la Bitácora es:\n"
        "1. PENDIENTE_REVISION: Creada por el empleado.\n"
        "2. COMENTADO_JEFE / REQUIERE_AJUSTE: El jefe responde solicitando cambios o agregando comentarios preliminares.\n"
        "3. CERRADO: El jefe aprueba y cierra la entrada con un comentario de feedback definitivo."
    )

    add_heading_1(doc, "6. Módulo 4: Bandejas de Autoevaluación y Evaluación")
    add_body_text(
        doc,
        "El líder tiene acceso a dos bandejas cruciales para el ciclo formal de desempeño."
    )
    
    add_heading_2(doc, "6.1 Autoevaluaciones")
    add_body_text(
        doc,
        "Permite ver las evaluaciones de desempeño que los propios empleados han realizado sobre sus objetivos "
        "y competencias soft skills, junto con sus comentarios, antes de proceder a la calificación de liderazgo."
    )
    add_screenshot(doc, "08_admin_autoevaluacion.png", "Bandeja de Autoevaluaciones del equipo completadas.")

    add_heading_2(doc, "6.2 Evaluaciones Formales")
    add_body_text(
        doc,
        "Panel donde el líder califica formalmente a su equipo. El flujo incluye la Revisión de Mitad de Año (Feedback Cuatrimestral) "
        "y la Evaluación Final Anual."
    )
    add_screenshot(doc, "09_admin_evaluacion.png", "Bandeja de Evaluaciones Formales del equipo con escala del 1 al 5.")

    create_callout_box(
        doc,
        "La escala de valoración oficial es de 1 a 5:\n"
        "1 - Malo | 2 - Regular | 3 - Bueno | 4 - Muy Bueno | 5 - Excelente.\n"
        "Los cálculos del rendimiento promedio del colaborador combinan los promedios matemáticos ponderados "
        "por el peso de cada objetivo."
    )

    add_heading_1(doc, "7. Módulo 5: Cursos, Capacitaciones y Calendario")
    add_heading_2(doc, "7.1 Panel de Cursos")
    add_body_text(
        doc,
        "Permite a los administradores registrar capacitaciones institucionales obligatorias u opcionales, "
        "y asignar dichos cursos a los empleados a cargo para realizar un seguimiento del cumplimiento de capacitación laboral."
    )
    add_screenshot(doc, "10_admin_cursos.png", "Catálogo de cursos asignados al personal de la organización.")

    add_heading_2(doc, "7.2 Calendario General")
    add_body_text(
        doc,
        "Muestra de forma unificada en una grilla mensual los hitos fundamentales del año, fechas límite de entrega de objetivos, "
        "reuniones programadas de mitad de año y periodos de autoevaluación."
    )
    add_screenshot(doc, "11_admin_calendario.png", "Calendario institucional mensual de hitos y deadlines.")

    add_heading_1(doc, "8. Módulo 6: Guía del Portal")
    add_body_text(
        doc,
        "La sección Guía expone de manera estática y configurable la documentación institucional sobre los pilares y "
        "competencias definidos por el departamento de Capital Humano de Permaquim S.A."
    )
    add_screenshot(doc, "12_admin_guia.png", "Vista de la Guía Institucional con pilares y competencias.")

    add_heading_1(doc, "9. Módulo 7: Administración y Control Global")
    add_body_text(
        doc,
        "Reservado exclusivamente para usuarios con rol administrador o superusuario (con la bandera EsSuperusuario activa)."
    )
    
    add_heading_2(doc, "9.1 Configuración de Plataforma")
    add_body_text(
        doc,
        "Permite administrar de forma dinámica los parámetros globales del sistema, tales como:\n"
        "- Email de Soporte (ayuda al usuario).\n"
        "- Días de próximo vencimiento (plazo de alerta).\n"
        "- Habilitación de objetivos específicos por área en evaluaciones.\n"
        "- Habilitación de cálculos de área comercial.\n"
        "- Posibilidad de autoevaluación del empleado (Activar/Desactivar módulo).\n"
        "- Un objetivo por pilar (restricción de unicidad)."
    )
    add_screenshot(doc, "13_admin_configuracion.png", "Panel de Configuración de parámetros globales del portal.")

    add_heading_2(doc, "9.2 Gestión de Usuarios (ABM)")
    add_body_text(
        doc,
        "Pantalla para gestionar la nómina de usuarios. El administrador puede dar de alta nuevos colaboradores o jefes, "
        "editar sus datos básicos, asociar jefes de reporte y restablecer la contraseña a la clave predeterminada de su legajo."
    )
    add_screenshot(doc, "14_admin_usuarios.png", "Grilla de Gestión de Usuarios y ABM de personal.")

    add_heading_1(doc, "10. Soporte y Contacto")
    add_body_text(
        doc,
        "Ante cualquier inconveniente, duda o consulta técnica sobre el funcionamiento del portal PQ-Talent, "
        "puede contactar directamente al área de Recursos Humanos de Permaquim enviando un correo electrónico a la casilla configurada: rrhh@permaquim.com."
    )

    doc.save("docs/Manual_Administrador.docx")
    print("Saved Manual_Administrador.docx")

def generate_employee_manual():
    doc = docx.Document()
    generate_cover_page(doc, "Manual de Usuario - Empleado", "Guía de Gestión del Rendimiento y Desarrollo Profesional para Colaboradores", "Antigravity AI Agent")
    add_header_footer(doc, "Manual de Empleado")
    
    add_heading_1(doc, "1. Introducción a PQ-Talent (Colaborador)")
    add_body_text(
        doc,
        "Bienvenido al manual del Colaborador para la plataforma PQ-Talent. "
        "Este portal le permite gestionar sus metas anuales, dar seguimiento diario a sus avances a través de la Bitácora, "
        "adjuntar evidencias de cumplimiento, chatear directamente con su supervisor y realizar sus procesos de autoevaluación formal."
    )
    
    add_heading_1(doc, "2. Inicio de Sesión y Primeros Pasos")
    add_body_text(
        doc,
        "Para acceder, ingrese su correo institucional y su legajo como contraseña inicial. "
        "Si el sistema detecta que es su primer ingreso, le solicitará de forma obligatoria cambiar su contraseña temporal."
    )
    add_screenshot(doc, "01_empleado_login.png", "Pantalla de Login del Colaborador.")

    create_callout_box(
        doc,
        "Su contraseña debe cumplir con los criterios de complejidad corporativos (mínimo 8 caracteres, al menos una letra y un número). "
        "Asegúrese de recordarla, de lo contrario deberá solicitar un reseteo al administrador de la plataforma."
    )

    add_heading_1(doc, "3. Módulo 1: Mi Dashboard Personal")
    add_body_text(
        doc,
        "El Dashboard personal consolida sus indicadores de avance. Presenta tarjetas de KPI (Total Objetivos, En Curso, "
        "Vencen Pronto, Pendientes de Revisión), un gráfico de dona que indica el progreso global de sus tareas y "
        "la lista de los próximos eventos programados."
    )
    add_screenshot(doc, "02_empleado_dashboard.png", "Vista general de Mi Dashboard Personal.")

    add_heading_1(doc, "4. Módulo 2: Mis Objetivos y Competencias")
    add_body_text(
        doc,
        "En este panel, el colaborador visualiza el listado detallado de sus metas asignadas para el año en curso. "
        "Puede observar a qué pilar corresponde, el peso porcentual asignado y el vencimiento de cada uno."
    )
    add_screenshot(doc, "03_empleado_objetivos.png", "Grilla de Mis Objetivos personales del año en curso.")

    create_callout_box(
        doc,
        "La suma total del peso de todos sus objetivos activos para un mismo año debe dar obligatoriamente el 100%. "
        "En caso de que no lo cumpla, el sistema le advertirá con un cartel en la parte superior. "
        "Coordine con su jefe directo (Pablo Tripodi) para ajustar los pesos si visualiza esta advertencia."
    )

    add_heading_1(doc, "5. Módulo 3: Detalle de Objetivo, Bitácora y Chat")
    add_body_text(
        doc,
        "Al ingresar a la vista de detalle de cualquiera de sus objetivos, podrá reportar sus avances diarios e interactuar con su jefe."
    )
    add_screenshot(doc, "04_empleado_objetivo_detalle.png", "Detalle de Objetivo con bitácora personal y chat con el supervisor.")

    add_heading_2(doc, "5.1 Registro de Bitácora de Avances")
    add_body_text(
        doc,
        "Haciendo clic en el botón 'Nueva Entrada' de la sección de Bitácora, puede describir hitos alcanzados, "
        "guardar reportes y cargar enlaces a documentos de evidencia. Su jefe recibirá una notificación para revisar "
        "su carga y brindarle feedback."
    )

    create_callout_box(
        doc,
        "Mantenga su bitácora actualizada regularmente. Sirve como base empírica de sus logros "
        "para las revisiones formales de mitad de año y de cierre."
    )

    add_heading_2(doc, "5.2 Chat Bidireccional de Feedback")
    add_body_text(
        doc,
        "En el panel derecho dispone de un chat directo integrado. "
        "Use esta vía rápida para resolver dudas de la meta, coordinar revisiones o avisar sobre demoras en la ejecución."
    )

    add_heading_1(doc, "6. Módulo 4: Autoevaluación de Desempeño")
    add_body_text(
        doc,
        "Cuando el departamento de Capital Humano inicie el periodo de revisiones, deberá ingresar al panel de 'Autoevaluación'. "
        "Allí completará el formulario indicando el nivel de cumplimiento percibido en cada objetivo y en las competencias soft skills vinculadas, "
        "ingresando comentarios justificativos y el enlace a evidencias."
    )
    add_screenshot(doc, "05_empleado_autoevaluacion.png", "Bandeja y Formulario de Autoevaluación Personal.")

    add_heading_1(doc, "7. Módulo 5: Mis Cursos y Capacitaciones")
    add_body_text(
        doc,
        "En la sección de 'Cursos', visualice los programas y capacitaciones institucionales que le han sido asignados. "
        "Se diferencia claramente entre cursos obligatorios (requeridos para su puesto) y sugeridos/opcionales."
    )
    add_screenshot(doc, "06_empleado_cursos.png", "Mesa de Cursos asignados y estado de aprobación.")

    add_heading_1(doc, "8. Módulo 6: Calendario de Vencimientos")
    add_body_text(
        doc,
        "Acceda al Calendario para ver en formato mensual sus propios plazos clave. "
        "Permite ver los deadlines específicos de entrega de avances y reuniones individuales de feedback agendadas con su líder."
    )
    add_screenshot(doc, "07_empleado_calendario.png", "Vista de Calendario Personal con plazos del colaborador.")

    add_heading_1(doc, "9. Módulo 7: Guía y Conceptos Clave")
    add_body_text(
        doc,
        "Sección estática donde puede consultar en cualquier momento la definición formal de los pilares de Permaquim, "
        "así como el diccionario conceptual de soft skills corporativas."
    )
    add_screenshot(doc, "08_empleado_guia.png", "Guía corporativa de pilares estratégicos y competencias blandas.")

    add_heading_1(doc, "10. Soporte y Ayuda")
    add_body_text(
        doc,
        "Para realizar consultas, informar fallos o sugerir mejoras en el portal, "
        "póngase en contacto con el equipo de Capital Humano enviando un correo a: rrhh@permaquim.com."
    )

    doc.save("docs/Manual_Empleado.docx")
    print("Saved Manual_Empleado.docx")

if __name__ == "__main__":
    main()
